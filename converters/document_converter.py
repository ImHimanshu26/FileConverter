import os
import tempfile
from pathlib import Path
import io

try:
    from docx import Document
    from docx2pdf import convert as docx_to_pdf
except ImportError:
    pass

try:
    import PyPDF2
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except ImportError:
    pass

class DocumentConverter:
    """Handles document format conversions"""
    
    def __init__(self):
        self.supported_inputs = ['pdf', 'doc', 'docx', 'txt', 'rtf']
        self.supported_outputs = ['pdf', 'docx', 'txt']
    
    def convert(self, input_path, output_format, original_filename):
        """Convert document to specified format"""
        
        input_ext = Path(input_path).suffix.lower()[1:]
        output_filename = self._get_output_filename(original_filename, output_format)
        
        # Read input file content
        if input_ext == 'txt':
            return self._convert_from_txt(input_path, output_format, output_filename)
        elif input_ext == 'pdf':
            return self._convert_from_pdf(input_path, output_format, output_filename)
        elif input_ext in ['doc', 'docx']:
            return self._convert_from_docx(input_path, output_format, output_filename)
        else:
            raise ValueError(f"Unsupported input format: {input_ext}")
    
    def _convert_from_txt(self, input_path, output_format, output_filename):
        """Convert from TXT format"""
        
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        if output_format == 'txt':
            return content.encode('utf-8'), output_filename, 'text/plain'
        
        elif output_format == 'pdf':
            return self._txt_to_pdf(content, output_filename)
        
        elif output_format == 'docx':
            return self._txt_to_docx(content, output_filename)
        
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _convert_from_pdf(self, input_path, output_format, output_filename):
        """Convert from PDF format"""
        
        if output_format == 'pdf':
            with open(input_path, 'rb') as f:
                return f.read(), output_filename, 'application/pdf'
        
        # Extract text from PDF
        text_content = self._extract_text_from_pdf(input_path)
        
        if output_format == 'txt':
            return text_content.encode('utf-8'), output_filename, 'text/plain'
        
        elif output_format == 'docx':
            return self._txt_to_docx(text_content, output_filename)
        
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _convert_from_docx(self, input_path, output_format, output_filename):
        """Convert from DOCX format"""
        
        if output_format == 'docx':
            with open(input_path, 'rb') as f:
                return f.read(), output_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Extract text from DOCX
        text_content = self._extract_text_from_docx(input_path)
        
        if output_format == 'txt':
            return text_content.encode('utf-8'), output_filename, 'text/plain'
        
        elif output_format == 'pdf':
            return self._txt_to_pdf(text_content, output_filename)
        
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _extract_text_from_pdf(self, pdf_path):
        """Extract text content from PDF"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception:
            return "Unable to extract text from PDF"
    
    def _extract_text_from_docx(self, docx_path):
        """Extract text content from DOCX"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception:
            return "Unable to extract text from DOCX"
    
    def _txt_to_pdf(self, text_content, output_filename):
        """Convert text to PDF"""
        try:
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            
            # Set up text formatting
            width, height = letter
            margin = 50
            line_height = 14
            lines_per_page = int((height - 2 * margin) / line_height)
            
            # Split text into lines
            lines = text_content.split('\n')
            
            y_position = height - margin
            page_num = 1
            
            for line in lines:
                # Wrap long lines
                words = line.split(' ')
                current_line = ""
                
                for word in words:
                    test_line = current_line + word + " "
                    if c.stringWidth(test_line) < (width - 2 * margin):
                        current_line = test_line
                    else:
                        if current_line:
                            c.drawString(margin, y_position, current_line.strip())
                            y_position -= line_height
                            
                            if y_position < margin:
                                c.showPage()
                                y_position = height - margin
                                page_num += 1
                        
                        current_line = word + " "
                
                if current_line:
                    c.drawString(margin, y_position, current_line.strip())
                    y_position -= line_height
                    
                    if y_position < margin:
                        c.showPage()
                        y_position = height - margin
                        page_num += 1
            
            c.save()
            buffer.seek(0)
            return buffer.getvalue(), output_filename, 'application/pdf'
            
        except Exception as e:
            raise Exception(f"Failed to convert to PDF: {str(e)}")
    
    def _txt_to_docx(self, text_content, output_filename):
        """Convert text to DOCX"""
        try:
            doc = Document()
            
            # Split text into paragraphs
            paragraphs = text_content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text)
                else:
                    doc.add_paragraph("")  # Empty paragraph for spacing
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer.getvalue(), output_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        except Exception as e:
            raise Exception(f"Failed to convert to DOCX: {str(e)}")
    
    def _get_output_filename(self, original_filename, output_format):
        """Generate output filename with new extension"""
        base_name = Path(original_filename).stem
        return f"{base_name}.{output_format}"
